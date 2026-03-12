"""
agents/cam_agent.py
Phase 5 – CAM Generator (Five Cs of Credit structure)

Generates a professional Credit Appraisal Memo (.docx) using the
Five Cs of Credit framework:
  C1 – Character   (Promoter profile, CIBIL, litigation)
  C2 – Capacity    (Revenue, EBITDA, ICR, DSCR)
  C3 – Capital     (Net Worth, D/E ratio)
  C4 – Collateral  (Asset type, value, charge type)
  C5 – Conditions  (Sector outlook, RBI regulations, macro)
"""

import io
import os
import logging
from datetime import date
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils.helpers import format_currency, risk_label

logger = logging.getLogger(__name__)

TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__), "..", "templates", "template.docx"
)


# ──────────────────────────────────────────────
# Build template.docx (Five Cs structure)
# ──────────────────────────────────────────────

def create_template() -> None:
    """Creates templates/template.docx with Five Cs structure and placeholders."""
    doc = Document()

    # ── Title ──
    title = doc.add_heading("CREDIT APPRAISAL MEMORANDUM (CAM)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ── Header Info Table ──
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    header_cells = [
        ("Company Name",      "{{company_name}}"),
        ("GSTIN / PAN",       "{{gstin}} / {{pan}}"),
        ("Loan Amount",       "{{loan_amount}}"),
        ("Loan Purpose",      "{{loan_purpose}}"),
        ("Date of Appraisal", "{{date}}"),
        ("Credit Score",      "{{credit_score}} / 100  —  {{risk_category}}"),
    ]
    for i, (lbl, ph) in enumerate(header_cells):
        table.cell(i, 0).text = lbl
        table.cell(i, 1).text = ph

    doc.add_paragraph()

    # ── VERDICT BOX ──
    doc.add_heading("RECOMMENDATION VERDICT", level=1)
    doc.add_paragraph("{{verdict}}")
    doc.add_paragraph("Approved Limit: {{approved_limit}}   |   Pricing: {{rate_label}}")
    doc.add_paragraph("{{verdict_reason}}")

    doc.add_paragraph()

    # ── C1: Character ──
    doc.add_heading("C1 – Character: Promoter Profile & Legal Standing", level=1)
    doc.add_paragraph("{{c1_character}}")

    # ── C2: Capacity ──
    doc.add_heading("C2 – Capacity: Repayment & Earnings Ability", level=1)
    doc.add_paragraph("{{c2_capacity}}")

    # ── C3: Capital ──
    doc.add_heading("C3 – Capital: Net Worth & Leverage", level=1)
    doc.add_paragraph("{{c3_capital}}")

    # ── C4: Collateral ──
    doc.add_heading("C4 – Collateral: Security & Charge", level=1)
    doc.add_paragraph("{{c4_collateral}}")

    # ── C5: Conditions ──
    doc.add_heading("C5 – Conditions: Market & Regulatory Environment", level=1)
    doc.add_paragraph("{{c5_conditions}}")

    # ── GST & Compliance ──
    doc.add_heading("6. GST Compliance & Bank Statement Analysis", level=1)
    doc.add_paragraph("{{gst_summary}}")
    doc.add_paragraph("Bank Statement Check: {{bank_check_summary}}")

    # ── Risk Summary ──
    doc.add_heading("7. Identified Risk Factors", level=1)
    doc.add_paragraph("{{risk_factors}}")

    # ── Signatures ──
    doc.add_paragraph()
    doc.add_paragraph("Prepared by: _________________________     Date: {{date}}")
    doc.add_paragraph("Reviewed by: _________________________")
    doc.add_paragraph("Approved by: _________________________")

    os.makedirs(os.path.dirname(TEMPLATE_PATH), exist_ok=True)
    doc.save(TEMPLATE_PATH)
    logger.info("Template saved to %s", TEMPLATE_PATH)


# ──────────────────────────────────────────────
# Placeholder replacement helpers
# ──────────────────────────────────────────────

def _replace_in_paragraph(para, replacements: dict) -> None:
    for key, value in replacements.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in para.text:
            for run in para.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(value))


def _replace_in_table(table, replacements: dict) -> None:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                _replace_in_paragraph(para, replacements)


def generate_cam(data: dict) -> bytes:
    """
    Fill the template with extracted + computed data.

    Args:
        data: dict from build_cam_data()

    Returns:
        bytes of the completed .docx file.
    """
    if not os.path.exists(TEMPLATE_PATH):
        logger.warning("template.docx not found – creating it now.")
        create_template()

    doc = Document(TEMPLATE_PATH)

    replacements = {
        "company_name":       data.get("company_name", "N/A"),
        "gstin":              data.get("gstin", "N/A"),
        "pan":                data.get("pan", "N/A"),
        "loan_amount":        format_currency(data.get("loan_amount", 0)),
        "loan_purpose":       data.get("loan_purpose", "N/A"),
        "credit_score":       data.get("credit_score", 0),
        "risk_category":      data.get("risk_category", risk_label(data.get("credit_score", 0))),
        "date":               data.get("date", date.today().strftime("%d %B %Y")),
        # Verdict
        "verdict":            data.get("verdict", "—"),
        "approved_limit":     format_currency(data.get("approved_limit", 0)),
        "rate_label":         data.get("rate_label", "—"),
        "verdict_reason":     data.get("verdict_reason", ""),
        # Five Cs
        "c1_character":       data.get("c1_character", ""),
        "c2_capacity":        data.get("c2_capacity", ""),
        "c3_capital":         data.get("c3_capital", ""),
        "c4_collateral":      data.get("c4_collateral", "Collateral details not provided."),
        "c5_conditions":      data.get("c5_conditions", ""),
        # Supporting
        "gst_summary":        data.get("gst_summary", "GST data not available."),
        "bank_check_summary": data.get("bank_check_summary", "Bank statement not provided."),
        "risk_factors":       data.get("risk_factors_text", "No specific risk factors identified."),
    }

    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    for table in doc.tables:
        _replace_in_table(table, replacements)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ──────────────────────────────────────────────
# Build cam_data from session results
# ──────────────────────────────────────────────

def build_cam_data(
    financials: dict,
    score_result: dict,
    gst_result: dict,
    sentiment_result: dict,
    score_explanation: str,
    collateral: str = "Not specified",
    bank_check: dict = None,
    verdict_data: dict = None,
) -> dict:
    """
    Consolidates all agent outputs into a single dict for generate_cam().

    Parameters
    ----------
    financials       : extracted financial KPIs
    score_result     : output from scoring_agent.compute_score()
    gst_result       : output from gst_agent.compare_gst()
    sentiment_result : output from web_research_agent
    score_explanation: LLM or rule-based recommendation text
    collateral       : credit officer input describing collateral
    bank_check       : output from bank_parser.check_circular_trading()
    verdict_data     : output from scoring_agent.compute_verdict()
    """
    fin  = financials or {}
    gst  = gst_result or {}
    sent = sentiment_result or {}
    verb = verdict_data or {}
    bank = bank_check or {}

    credit_score = score_result.get("total_score", 0)
    promoters    = fin.get("promoters", [])
    promoter_str = "; ".join(
        f"{p.get('name', 'N/A')} ({p.get('designation', '')})"
        for p in promoters
    ) if promoters else "Promoter details not extracted."

    # C1 – Character
    c1 = (
        f"Promoters: {promoter_str}\n"
        f"CIBIL Score: {score_result.get('sub_scores', {}).get('legal_cibil', 0) * 6 + 300:.0f} (mapped)\n"
        f"News Sentiment: {sent.get('overall_sentiment', 'N/A').capitalize()}\n"
        f"Key Risks from Secondary Research: "
        + "; ".join(sent.get("key_risks", ["Not available"]))
    )

    # C2 – Capacity
    c2 = (
        f"Total Revenue: {format_currency(fin.get('total_revenue', 0))}\n"
        f"Net Profit: {format_currency(fin.get('net_profit', 0))}\n"
        f"EBITDA: {format_currency(fin.get('ebitda', 0))}\n"
        f"Interest Coverage Ratio: {fin.get('interest_coverage_ratio') or fin.get('interest_coverage', 'N/A')}\n"
        f"Current Ratio: {fin.get('current_ratio', 'N/A')}\n"
        f"Loan Amount Requested: {format_currency(fin.get('loan_amount_requested', 0))}"
    )

    # C3 – Capital
    c3 = (
        f"Net Worth: {format_currency(fin.get('net_worth', 0))}\n"
        f"Total Assets: {format_currency(fin.get('total_assets', 0))}\n"
        f"Total Loans / Debt: {format_currency(fin.get('total_loans', 0))}\n"
        f"Debt / Equity Ratio: {fin.get('debt_equity_ratio', 'N/A')}"
    )

    # C4 – Collateral
    c4 = collateral if collateral and collateral.strip() else "Collateral details not specified by credit officer."

    # C5 – Conditions
    sent_summary = sent.get("summary", "Market conditions data not available.")
    positives    = "; ".join(sent.get("key_positives", [])) or "None noted."
    c5 = (
        f"Market Sentiment: {sent.get('overall_sentiment', 'neutral').capitalize()}\n"
        f"Sector Positives: {positives}\n"
        f"Analyst Summary: {sent_summary}"
    )

    # GST summary
    gst_para = gst.get("flag_message", "GST data not available.")

    # Bank check summary
    if bank:
        bank_para = bank.get("flag_message", "Bank statement check not available.")
    else:
        bank_para = "Bank statement not uploaded — check not performed."

    # Risk factors
    raw_risks = fin.get("risk_factors", [])
    risk_text = "\n".join(f"  • {r}" for r in raw_risks) if raw_risks else "No specific risk factors identified."

    return {
        "company_name":       fin.get("company_name", "Borrower Company"),
        "gstin":              fin.get("gstin", "N/A"),
        "pan":                fin.get("pan", "N/A"),
        "loan_amount":        fin.get("loan_amount_requested", 0),
        "loan_purpose":       fin.get("loan_purpose", "Working Capital"),
        "credit_score":       credit_score,
        "risk_category":      risk_label(credit_score),
        "date":               date.today().strftime("%d %B %Y"),
        # Verdict
        "verdict":            verb.get("verdict", "—"),
        "approved_limit":     verb.get("approved_limit", 0),
        "rate_label":         verb.get("rate_label", "—"),
        "verdict_reason":     verb.get("reason", score_explanation),
        # Five Cs
        "c1_character":       c1,
        "c2_capacity":        c2,
        "c3_capital":         c3,
        "c4_collateral":      c4,
        "c5_conditions":      c5,
        # Supporting sections
        "gst_summary":        gst_para,
        "bank_check_summary": bank_para,
        "risk_factors_text":  risk_text,
    }
